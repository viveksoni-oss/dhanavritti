from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Dhanavritti_Website_Client_Brief.docx"

GREEN = RGBColor(8, 96, 32)
DARK = RGBColor(26, 26, 26)
MUTED = RGBColor(84, 98, 87)
LIGHT_GREEN = "F0F9F3"
PALE = "F7FBF4"
BORDER = "D9EAD3"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_run(run, size=11, bold=False, color=DARK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_rich_para(doc, parts, after=6, before=0, line_spacing=1.1, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    for text, bold in parts:
        run = p.add_run(text)
        style_run(run, bold=bold, color=DARK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 11)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=16 if level == 1 else 13, bold=True, color=GREEN)
    return p


def add_bullet(doc, parts):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    for text, bold in parts:
        run = p.add_run(text)
        style_run(run, bold=bold)
    return p


def add_number(doc, number, parts):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    marker = p.add_run(f"{number}. ")
    style_run(marker, bold=True, color=GREEN)
    for text, bold in parts:
        run = p.add_run(text)
        style_run(run, bold=bold)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, LIGHT_GREEN)
    set_cell_border(cell, color=BORDER)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    style_run(r, size=12, bold=True, color=GREEN)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    style_run(r2, size=10.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_flow_table(doc, title, steps):
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=len(steps), cols=2)
    set_table_width(table, [Inches(1.0), Inches(5.35)])
    for idx, (label, detail) in enumerate(steps):
        left = table.cell(idx, 0)
        right = table.cell(idx, 1)
        for cell in (left, right):
            set_cell_border(cell, color=BORDER)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, PALE)
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        style_run(r, size=10, bold=True, color=GREEN)
        p2 = right.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(detail)
        style_run(r2, size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1


def build_doc():
    doc = Document()
    configure_document(doc)

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("DHANAVRITTI WEBSITE")
    style_run(r, size=11, bold=True, color=GREEN)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Client Brief: Website Value, Technology & Lead Flow")
    style_run(r, size=22, bold=True, color=DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run(
        "A modern, secure, high-performance website built for credibility, investor confidence, and structured application handling."
    )
    style_run(r, size=11, color=MUTED)

    add_callout(
        doc,
        "Executive Summary",
        "The website is not a basic WordPress CMS build. It is a custom Next.js application with a premium interface, fast deployment pipeline, secure form handling, branded transactional emails, CAPTCHA protection, and a future-ready backend architecture.",
    )

    add_heading(doc, "What Makes This Website Strong", level=1)
    bullets = [
        [("Modern custom technology stack: ", True), ("Built using Next.js, React, Tailwind CSS, and Vercel instead of a heavy WordPress theme/plugin stack.", False)],
        [("Premium brand experience: ", True), ("Custom visual design, refined spacing, brand-led colors, polished sections, and a professional investment-platform feel.", False)],
        [("Fast and smooth user experience: ", True), ("No unnecessary WordPress plugins or page-builder bloat, resulting in faster loading and a cleaner browsing experience.", False)],
        [("High-quality animation layer: ", True), ("Smooth, controlled animations create a premium feel without making the site feel slow or heavy.", False)],
        [("Secure application handling: ", True), ("Server-side validation, file type checks, file size checks, CAPTCHA protection, and rate limiting help protect against spam and misuse.", False)],
        [("Professional email workflow: ", True), ("The platform sends branded HTML emails to the internal team and confirmation emails to applicants.", False)],
        [("Better email deliverability foundation: ", True), ("Resend can be used with a verified sending domain, SPF, DKIM, and DMARC records to reduce spam risk and improve transactional email reliability.", False)],
        [("Deployment pipeline ready: ", True), ("Future updates can be deployed quickly through the configured Vercel pipeline instead of manual CMS edits.", False)],
        [("Future-ready architecture: ", True), ("The site can later support database storage, Google Sheets sync, dashboards, analytics, CRM-style lead tracking, and admin workflows.", False)],
    ]
    for item in bullets:
        add_bullet(doc, item)

    add_heading(doc, "Why This Is Better Than A Typical WordPress CMS Website", level=1)
    comparison = doc.add_table(rows=1, cols=3)
    set_table_width(comparison, [Inches(1.65), Inches(2.35), Inches(2.5)])
    headers = ["Area", "Typical WordPress CMS", "Current Custom Website"]
    for i, h in enumerate(headers):
        cell = comparison.cell(0, i)
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell)
        set_cell_margins(cell)
        r = cell.paragraphs[0].add_run(h)
        style_run(r, size=10.5, bold=True, color=GREEN)
    rows = [
        ("Performance", "Often slowed by themes, plugins, and page builders.", "Lightweight custom frontend with faster page experience."),
        ("Design", "Commonly template-based and less differentiated.", "Premium custom look designed around the Dhanavritti brand."),
        ("Security", "Plugin/admin vulnerabilities must be continuously managed.", "Smaller attack surface with controlled custom code."),
        ("Forms", "Usually dependent on third-party form plugins.", "Custom application API with validation, file checks, CAPTCHA, and rate limiting."),
        ("Email", "Often depends on hosting mail or SMTP plugins.", "Transactional email workflow using Resend-ready infrastructure."),
        ("Scalability", "Can become messy as plugins increase.", "Can evolve into backend database, Sheets sync, dashboard, and CRM-like workflows."),
    ]
    for area, wp, current in rows:
        row = comparison.add_row().cells
        values = [area, wp, current]
        for i, value in enumerate(values):
            cell = row[i]
            set_cell_border(cell)
            set_cell_margins(cell)
            if i == 0:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            style_run(r, size=9.6, bold=(i == 0), color=GREEN if i == 0 else DARK)

    add_heading(doc, "Current Application Flow", level=1)
    add_flow_table(
        doc,
        "Current Flow",
        [
            ("Step 1", "Founder visits the Apply page and fills company, founder, email, phone, and pitch deck details."),
            ("Step 2", "Cloudflare Turnstile CAPTCHA and server-side checks protect the form from bots and spam submissions."),
            ("Step 3", "Uploaded files are validated by size and allowed file type before processing."),
            ("Step 4", "The internal Dhanavritti team receives a branded application email with the applicant details and attached pitch deck."),
            ("Step 5", "The applicant receives a branded confirmation email acknowledging that the application has been received."),
            ("Step 6", "If backend database credentials are enabled, application metadata can also be stored in the backend database."),
        ],
    )

    add_heading(doc, "Recommended Future Enhancement Flow", level=1)
    add_rich_para(
        doc,
        [
            ("If the client wants a more structured lead-management process, the website can be extended with ", False),
            ("Google Sheet sync, secure pitch deck links, and backend database storage", True),
            (". This can be treated as a separate enhancement phase after approval.", False),
        ],
        after=8,
    )
    add_flow_table(
        doc,
        "Future Flow With Sheet + Backend Database",
        [
            ("Step 1", "Founder submits the application form and uploads the pitch deck."),
            ("Step 2", "Pitch deck is saved to secure storage and a private/shareable file link is generated."),
            ("Step 3", "Application details are stored in the backend database for long-term tracking and future dashboard use."),
            ("Step 4", "A Google Sheet is automatically updated with company name, founder name, email, phone, timestamp, status, and pitch deck link."),
            ("Step 5", "The Dhanavritti team receives a branded email notification with the key details and deck access link."),
            ("Step 6", "The applicant receives a confirmation email, while the team can later search, filter, export, or build review workflows from the stored data."),
        ],
    )

    add_heading(doc, "Important Notes For The Client", level=1)
    notes = [
        [("Email deliverability: ", True), ("No provider can honestly guarantee that emails will never enter spam. However, using Resend with verified domain records such as SPF, DKIM, and DMARC gives a much stronger professional deliverability setup than basic hosting email.", False)],
        [("Official recipient email: ", True), ("The current setup can send applications to the official Dhanavritti contact email once the final email address and domain settings are confirmed.", False)],
        [("Database and Sheet sync: ", True), ("Database storage and Google Sheet automation are valuable future add-ons if the team wants structured lead tracking beyond email notifications.", False)],
        [("Maintenance advantage: ", True), ("The custom setup avoids common WordPress plugin conflicts and allows cleaner future development.", False)],
    ]
    for item in notes:
        add_bullet(doc, item)

    add_heading(doc, "Suggested Client Positioning", level=1)
    add_callout(
        doc,
        "Positioning Statement",
        "This website is built as a modern digital platform for Dhanavritti, not just as an informational website. It supports brand credibility, a premium first impression, secure application intake, professional communication, and future business-process automation.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
